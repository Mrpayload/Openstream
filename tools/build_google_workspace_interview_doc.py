from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"V:\Antgravity\webstreamer\interview-docs\Google_Workspace_Zero_to_Hero_Interview_Answers.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullet(doc, text, style="List Bullet"):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_answer_box(doc, title, answer):
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    header = table.cell(0, 0)
    body = table.cell(1, 0)
    shade_cell(header, "1F4E79")
    set_cell_text(header, title, bold=True, color="FFFFFF")
    set_cell_text(body, answer)
    for p in body.paragraphs:
        for run in p.runs:
            run.font.size = Pt(10.5)
    doc.add_paragraph()


def set_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 24, "1F4E79"),
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "2F5597"),
        ("Heading 3", 11.5, "404040"),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos Display" if style_name == "Title" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    set_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Google Workspace Administration").bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Zero to Hero Interview Answers in Plain Language")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph(
        "Use this document to speak confidently about Google Workspace support in an IT Support or IT Operations interview. "
        "The goal is not to sound like a textbook. The goal is to explain what you can do, how you troubleshoot, and why access management matters in a real business."
    )

    doc.add_heading("1. Simple Version of the Skill", level=1)
    doc.add_paragraph(
        "Google Workspace administration means helping a business manage Google services such as Gmail, Drive, Meet, Calendar, Admin Console, and Chromebooks. "
        "In simple terms, it is about making sure users can access the tools they need, company data stays protected, and common issues are resolved quickly."
    )

    add_answer_box(
        doc,
        "Short Interview Answer",
        "I have experience supporting Google Workspace services such as Gmail, Drive, Meet, Calendar, Admin Console, and Chromebook management. "
        "I can help users with account access, shared drives, email issues, permissions, and device-related problems. "
        "I also understand that managing access correctly is important when staff join, change roles, or leave the company."
    )

    doc.add_heading("2. Stronger Interview Answer", level=1)
    add_answer_box(
        doc,
        "Stronger Answer",
        "In my Google Workspace support experience, I focus on both user support and secure administration. "
        "I can help with Gmail, Drive, Meet, Calendar, Admin Console, shared drives, permissions, and Chromebook-related issues. "
        "For example, if a user cannot access a shared file, I check whether the problem is with their account, group membership, file permissions, or shared drive settings. "
        "I also support onboarding and offboarding by making sure new users get the right access and leaving users lose access at the right time. "
        "My approach is to resolve the immediate issue, protect company data, and document anything that may help the team in the future."
    )

    doc.add_heading("3. Zero to Hero Breakdown", level=1)

    sections = [
        (
            "Gmail Support",
            [
                "Help users with login problems, mailbox access, sending and receiving email, spam, filters, forwarding, and email client setup.",
                "Check whether the issue affects one user, multiple users, or the whole organisation.",
                "Use Admin Console checks when needed, such as account status, aliases, routing, and security settings.",
            ],
            "If a user says they are not receiving emails, I first check whether the issue is with one sender, all senders, or a specific mailbox rule. Then I check spam, filters, forwarding, storage, account status, and any admin-level mail routing or security settings.",
        ),
        (
            "Google Drive and Shared Drives",
            [
                "Support file access, ownership, shared folders, shared drives, external sharing, and deleted or missing files.",
                "Understand the difference between My Drive and Shared Drives.",
                "Make sure access is based on business need, not convenience only.",
            ],
            "With Drive issues, I usually check the user account, the file location, the sharing settings, and whether access is granted directly or through a group. For Shared Drives, I also check the member role, because Viewer, Contributor, Content Manager, and Manager permissions behave differently.",
        ),
        (
            "Google Meet and Calendar",
            [
                "Support meeting access, calendar invitations, room/resource booking, video or microphone issues, and permission problems.",
                "Check browser, account, network, and device settings when meetings fail.",
                "Help users understand calendar sharing and visibility settings.",
            ],
            "For Meet and Calendar issues, I check whether the problem is account-related, browser-related, device-related, or permission-related. For example, if someone cannot join a meeting, I check the invitation, account being used, browser permissions, camera/microphone settings, and network connectivity.",
        ),
        (
            "Admin Console",
            [
                "Use Admin Console to manage users, groups, organisational units, devices, security settings, and service access.",
                "Check account status, password resets, 2-step verification, group membership, and app access.",
                "Apply least-privilege access so users receive only the access they need.",
            ],
            "The Admin Console is where I would manage and troubleshoot user access. I use it to check whether the account is active, whether the user is in the correct group or organisational unit, and whether any security policy or service restriction is blocking access.",
        ),
        (
            "Chromebook Management",
            [
                "Support Chromebook enrolment, sign-in, Wi-Fi, updates, profiles, and basic device troubleshooting.",
                "Understand that Chromebooks can be controlled through Google Admin policies.",
                "Escalate hardware issues when needed while still checking configuration first.",
            ],
            "For Chromebook issues, I check whether the device is enrolled, whether the user is signing in with the correct account, whether Wi-Fi is working, and whether any device or user policy is causing the problem. I also check for updates and basic hardware symptoms before escalating.",
        ),
        (
            "User Access Management",
            [
                "Create, update, suspend, and remove user access during onboarding, role changes, and offboarding.",
                "Use groups and organisational units to manage access more cleanly.",
                "Protect business data by removing access quickly when staff leave.",
            ],
            "I see access management as one of the most important parts of Google Workspace administration. When staff join, they need the right access to work from day one. When staff change roles, access should be updated. When staff leave, accounts and data access must be secured quickly to reduce risk.",
        ),
    ]

    for heading, bullets, answer in sections:
        doc.add_heading(heading, level=2)
        for bullet in bullets:
            add_bullet(doc, bullet)
        add_answer_box(doc, "Interview Answer", answer)

    doc.add_heading("4. Common Interview Questions and Answers", level=1)
    qas = [
        (
            "How would you troubleshoot a user who cannot access a Google Drive file?",
            "I would first confirm which account the user is signed into and whether they are using the correct file link. Then I would check the file or folder permissions, shared drive membership, group membership, and whether external sharing restrictions apply. If access should be granted, I would follow the company approval process before changing permissions."
        ),
        (
            "How do you handle a new employee in Google Workspace?",
            "I would create or activate their account, place them in the correct organisational unit or groups, assign the required services, help with initial sign-in and 2-step verification, and confirm they can access Gmail, Drive, Calendar, and any shared resources needed for their role."
        ),
        (
            "How do you handle an employee leaving the company?",
            "I would follow the offboarding checklist. That usually means suspending or securing the account, resetting active sessions if needed, transferring important Drive ownership or shared drive access, removing group membership, recovering devices, and documenting the completion of the process."
        ),
        (
            "What would you do if multiple users report Gmail issues?",
            "I would check whether it is a wider service issue or an internal configuration issue. I would review the Google Workspace Status Dashboard, confirm the scope of impact, check recent admin changes, and communicate updates to users while troubleshooting."
        ),
        (
            "How do you manage permissions safely?",
            "I follow least privilege. I try to give users only the access they need for their job. I prefer group-based access where possible because it is easier to manage and audit. I also avoid making files public unless there is a clear business reason and approval."
        ),
    ]

    for question, answer in qas:
        doc.add_heading(question, level=2)
        doc.add_paragraph(answer)

    doc.add_heading("5. STAR Example Answer", level=1)
    doc.add_paragraph("Use this structure when the interviewer asks for a real example.")
    star = [
        ("Situation", "A user or team could not access important files or Google services needed for daily work."),
        ("Task", "My responsibility was to restore access quickly while making sure company data stayed secure."),
        ("Action", "I confirmed the user account, checked group membership, reviewed file or shared drive permissions, checked Admin Console settings, and applied the correct access after confirming business need."),
        ("Result", "The user regained access, the permission issue was resolved properly, and the fix was documented so similar issues could be handled faster next time."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "STAR Part", bold=True, color="FFFFFF")
    set_cell_text(table.rows[0].cells[1], "Plain-Language Example", bold=True, color="FFFFFF")
    shade_cell(table.rows[0].cells[0], "1F4E79")
    shade_cell(table.rows[0].cells[1], "1F4E79")
    for label, text in star:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], text)

    doc.add_heading("6. Words to Use in the Interview", level=1)
    useful_phrases = [
        "I usually start by checking the scope of the issue.",
        "I confirm whether it is user-specific, device-specific, permission-related, or service-wide.",
        "I follow least privilege when granting access.",
        "I prefer group-based access because it is easier to manage and audit.",
        "I make sure offboarding is completed properly because it directly affects security.",
        "I document repeated issues so the team can resolve them faster in the future.",
    ]
    for phrase in useful_phrases:
        add_bullet(doc, phrase)

    doc.add_heading("7. Final Polished Answer to Memorise", level=1)
    add_answer_box(
        doc,
        "Final Answer",
        "I have practical experience supporting Google Workspace services including Gmail, Drive, Meet, Calendar, Admin Console, and Chromebooks. "
        "I help users with login issues, email problems, file access, shared drive permissions, calendar and meeting issues, and basic device troubleshooting. "
        "When I troubleshoot, I try to identify whether the issue is with the user account, device, browser, permission, group membership, or a wider service issue. "
        "I also understand that access management is very important. During onboarding, users need the right access from day one. During role changes, access should be updated. During offboarding, accounts and company data must be secured quickly. "
        "My focus is always to solve the user problem, protect company data, and keep the process documented."
    )

    doc.add_section(WD_SECTION.CONTINUOUS)
    footer = doc.sections[-1].footer.paragraphs[0]
    footer.text = "Google Workspace Interview Preparation - Plain Language Guide"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
