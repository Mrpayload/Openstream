from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"V:\Antgravity\webstreamer\interview-docs\SaaS_Administration_Zero_to_Hero_Interview_Answers.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    paragraph.paragraph_format.space_after = Pt(3)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    return paragraph


def add_answer_box(doc, title, answer):
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    header = table.cell(0, 0)
    body = table.cell(1, 0)
    shade_cell(header, "1F4E79")
    set_cell_text(header, title, bold=True, color="FFFFFF")
    set_cell_text(body, answer)
    doc.add_paragraph()


def set_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 24, "1F4E79"),
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "2F5597"),
        ("Heading 3", 11.5, "404040"),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos Display" if style_name == "Title" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    set_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("SaaS Administration").bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Zero to Hero Interview Answers in Plain Language")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph(
        "Use this guide to explain SaaS administration clearly in an IT Support or IT Operations interview. "
        "SaaS administration is about managing cloud-based business tools so users have the right access, licences are used properly, renewals are controlled, and vendor issues are followed through to resolution."
    )

    doc.add_heading("1. Simple Version of the Skill", level=1)
    doc.add_paragraph(
        "SaaS stands for Software as a Service. These are online tools that a business pays for and uses through the cloud. "
        "Examples include Microsoft 365, Google Workspace, Slack, Zoom, Xero, Adobe, Jira, ServiceNow, HR systems, CRM systems, password managers, and other business applications. "
        "SaaS administration means managing users, licences, permissions, subscriptions, renewals, support requests, and vendor communication."
    )

    add_answer_box(
        doc,
        "Short Interview Answer",
        "I have managed SaaS tools by handling licences, subscriptions, renewals, access permissions, and vendor coordination. "
        "I keep track of who has access to which systems and make sure licences are used properly. "
        "I also help with renewals and support requests by communicating with vendors and following up until the issue is resolved."
    )

    doc.add_heading("2. Stronger Interview Answer", level=1)
    add_answer_box(
        doc,
        "Stronger Answer",
        "I have practical experience supporting SaaS administration across user access, licence tracking, subscription renewals, and vendor coordination. "
        "My approach is to make sure users have the right access for their role, licences are not wasted, renewals are tracked before they become urgent, and vendor support requests are clearly communicated and followed up. "
        "I also understand the security side of SaaS administration. When users join, change roles, or leave, their access must be updated quickly so company systems and data stay protected."
    )

    doc.add_heading("3. Zero to Hero Breakdown", level=1)
    sections = [
        (
            "SaaS User Access",
            [
                "Create, update, disable, or remove user access in cloud-based systems.",
                "Assign the correct role or permission level based on the user's job.",
                "Review access during onboarding, role changes, and offboarding.",
            ],
            "For SaaS access, I focus on giving users the right access for their role. I check whether they need basic user access, admin access, manager access, or read-only access. I also make sure access is removed when someone leaves or no longer needs the tool.",
        ),
        (
            "Licence Management",
            [
                "Track how many licences are purchased, assigned, unused, or due for renewal.",
                "Remove licences from departed users or inactive accounts.",
                "Help the business avoid paying for unused or duplicated software.",
            ],
            "Licence management is important because SaaS costs can grow quickly. I track which users have licences, remove licences that are no longer needed, and help make sure the business only pays for what it actually uses.",
        ),
        (
            "Subscriptions and Renewals",
            [
                "Track renewal dates, contract terms, subscription owners, and vendor contacts.",
                "Give early notice before renewals so the business can review cost and usage.",
                "Help avoid last-minute renewals, service disruption, or surprise invoices.",
            ],
            "For renewals, I try to stay ahead of the date. I check usage, confirm whether the tool is still needed, review licence counts, and coordinate with the vendor or internal owner before the renewal deadline.",
        ),
        (
            "Permissions and Role-Based Access",
            [
                "Assign permissions based on job role and business need.",
                "Avoid unnecessary admin access.",
                "Use groups, teams, departments, or roles where the system supports them.",
            ],
            "I follow least privilege when managing SaaS permissions. That means users should get the access they need to do their job, but not extra admin rights unless there is a clear reason and approval.",
        ),
        (
            "Vendor Coordination",
            [
                "Contact vendors for technical issues, billing questions, renewals, and account problems.",
                "Provide clear information such as screenshots, error messages, affected users, and impact.",
                "Track vendor responses and follow up until the issue is resolved.",
            ],
            "When working with vendors, I try to give them clear information from the start. I explain the issue, business impact, affected users, screenshots or error messages, and what troubleshooting has already been done. Then I track the ticket and follow up until resolution.",
        ),
        (
            "Support Requests",
            [
                "Help users with login issues, access errors, application settings, licence problems, and feature questions.",
                "Separate user training issues from technical or vendor-side issues.",
                "Escalate to the vendor when the problem cannot be fixed internally.",
            ],
            "For SaaS support requests, I first check whether the issue is account-related, permission-related, licence-related, browser-related, or a wider service issue. If it needs vendor help, I collect the right details and raise a support ticket.",
        ),
        (
            "Security and Offboarding",
            [
                "Remove access from SaaS tools when staff leave.",
                "Transfer ownership of files, reports, workflows, or admin responsibilities where needed.",
                "Check for shared accounts, personal email access, or unused admin accounts.",
            ],
            "SaaS offboarding is a security priority. When someone leaves, I make sure their access is disabled or removed, licences are recovered, and any important ownership or data is transferred before the account is deleted or archived.",
        ),
        (
            "SaaS Documentation",
            [
                "Maintain records of application owners, vendor contacts, renewal dates, licence counts, and admin access.",
                "Document common support steps and escalation paths.",
                "Keep a simple register so the business knows what tools it uses and who manages them.",
            ],
            "Documentation makes SaaS administration easier because the team can quickly see what tools exist, who owns them, when they renew, who has admin access, and how to get support.",
        ),
    ]

    for heading, bullets, answer in sections:
        doc.add_heading(heading, level=2)
        for bullet in bullets:
            add_bullet(doc, bullet)
        add_answer_box(doc, "Interview Answer", answer)

    doc.add_heading("4. Common Interview Questions and Answers", level=1)
    qas = [
        (
            "What is SaaS administration?",
            "SaaS administration means managing cloud-based software tools used by a business. This includes user accounts, licences, permissions, subscriptions, renewals, vendor support, and security controls."
        ),
        (
            "How do you decide what access a user should get?",
            "I would check the user's role, department, manager approval, and business need. I follow least privilege, so the user gets the access required for their job without unnecessary admin rights."
        ),
        (
            "How do you manage unused licences?",
            "I would review assigned licences regularly, check inactive users or departed staff, remove unused licences, and update the licence register. Before renewals, I would compare actual usage against the number of licences being paid for."
        ),
        (
            "How do you handle a SaaS renewal?",
            "I would check the renewal date early, review usage and licence count, confirm whether the business still needs the tool, check if there are unused licences, and coordinate with the vendor or internal owner before approval."
        ),
        (
            "How do you work with vendors?",
            "I give vendors clear information: what the issue is, who is affected, business impact, screenshots, error messages, and steps already tried. I track the case number and follow up until the issue is resolved."
        ),
        (
            "What would you do if a user cannot access a SaaS application?",
            "I would check whether the account is active, whether the user has the right licence, whether permissions are correct, whether MFA or SSO is blocking access, and whether the application has a wider outage. If needed, I would escalate to the vendor."
        ),
        (
            "Why is SaaS offboarding important?",
            "It is important because many SaaS tools contain company data. If access is not removed when someone leaves, there is a security risk. Offboarding also helps recover licences and reduce unnecessary cost."
        ),
    ]

    for question, answer in qas:
        doc.add_heading(question, level=2)
        doc.add_paragraph(answer)

    doc.add_heading("5. STAR Example Answer", level=1)
    doc.add_paragraph("Use this when the interviewer asks for a real example of SaaS administration.")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    set_cell_text(headers[0], "STAR Part", bold=True, color="FFFFFF")
    set_cell_text(headers[1], "Plain-Language Example", bold=True, color="FFFFFF")
    shade_cell(headers[0], "1F4E79")
    shade_cell(headers[1], "1F4E79")
    rows = [
        ("Situation", "The business used multiple SaaS tools and needed better visibility of licences, users, renewals, and vendor contacts."),
        ("Task", "My responsibility was to help manage access, track licences, and support renewals and vendor issues."),
        ("Action", "I maintained records of users and licences, removed access for departed staff, checked renewal dates, coordinated with vendors, and followed up on support tickets."),
        ("Result", "The business had better control over SaaS access and cost, renewals were easier to manage, and user issues were resolved more consistently."),
    ]
    for label, text in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], text)

    doc.add_heading("6. Words to Use in the Interview", level=1)
    phrases = [
        "I track users, licences, subscriptions, renewals, and vendor contacts.",
        "I follow least privilege when assigning SaaS permissions.",
        "I remove licences from departed users or inactive accounts to reduce waste.",
        "I try to review renewals before they become urgent.",
        "I coordinate with vendors by providing clear details and following up until resolution.",
        "I treat SaaS offboarding as a security task, not just an admin task.",
        "I keep SaaS records updated so the business knows what tools it pays for and who has access.",
    ]
    for phrase in phrases:
        add_bullet(doc, phrase)

    doc.add_heading("7. Final Polished Answer to Memorise", level=1)
    add_answer_box(
        doc,
        "Final Answer",
        "I have experience supporting SaaS administration by managing user access, licences, subscriptions, renewals, permissions, and vendor coordination. "
        "I keep track of who has access to which systems, what licences are assigned, when renewals are due, and who to contact for vendor support. "
        "I also understand the security side of SaaS administration. Users should only have the access they need, admin access should be limited, and access must be removed quickly during offboarding. "
        "When issues come up, I troubleshoot the user account, licence, permission, SSO or MFA settings, and application status. If vendor support is needed, I provide clear information and follow up until the issue is resolved. "
        "Overall, I see SaaS administration as important for cost control, security, user productivity, and smooth IT operations."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "SaaS Administration Interview Preparation - Plain Language Guide"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
