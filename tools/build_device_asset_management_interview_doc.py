from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"V:\Antgravity\webstreamer\interview-docs\Device_and_Asset_Management_Zero_to_Hero_Interview_Answers.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_answer_box(doc, title, answer):
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    header = table.cell(0, 0)
    body = table.cell(1, 0)
    shade_cell(header, "1F4E79")
    set_cell_text(header, title, bold=True, color="FFFFFF")
    set_cell_text(body, answer)
    doc.add_paragraph()


def set_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 24, "1F4E79"),
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "2F5597"),
        ("Heading 3", 11.5, "404040"),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos Display" if style_name == "Title" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    set_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Device and Asset Management").bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Zero to Hero Interview Answers in Plain Language")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph(
        "Use this guide to explain device and asset management clearly in an IT Support or IT Operations interview. "
        "The main idea is simple: know what the company owns, who has it, where it is, what condition it is in, and what needs to happen next."
    )

    doc.add_heading("1. Simple Version of the Skill", level=1)
    doc.add_paragraph(
        "Device and asset management means keeping track of company IT equipment and software. "
        "This includes laptops, desktops, tablets, mobile phones, printers, monitors, accessories, licences, and business applications. "
        "A good asset register helps the IT team support users, reduce unnecessary spending, protect company data, and manage onboarding and offboarding smoothly."
    )

    add_answer_box(
        doc,
        "Short Interview Answer",
        "I have maintained hardware and software asset registers, including laptops, tablets, mobiles, printers, and other IT equipment. "
        "I track devices from purchase through deployment, recovery, repair, and replacement. "
        "Good asset management helps control cost, supports security, and makes onboarding and offboarding much smoother."
    )

    doc.add_heading("2. Stronger Interview Answer", level=1)
    add_answer_box(
        doc,
        "Stronger Answer",
        "I have practical experience managing IT assets across their full lifecycle. "
        "That includes recording new equipment when it is purchased, assigning it to users, tracking serial numbers and warranty details, updating the asset register when devices move, and making sure equipment is recovered during offboarding. "
        "I also track software licences and subscriptions so the business knows what it is paying for and who is using each tool. "
        "For me, asset management is not just admin work. It supports security, budgeting, compliance, onboarding, offboarding, and faster troubleshooting."
    )

    doc.add_heading("3. Zero to Hero Breakdown", level=1)
    sections = [
        (
            "Hardware Asset Register",
            [
                "Record laptops, desktops, tablets, mobiles, printers, monitors, docks, chargers, and other IT equipment.",
                "Track asset tag, serial number, model, assigned user, location, purchase date, warranty status, and condition.",
                "Keep the register updated whenever a device is issued, returned, repaired, replaced, or disposed of.",
            ],
            "A hardware asset register gives IT a clear view of what equipment the company owns and where it is. In an interview, I would explain that I track details like serial number, assigned user, location, warranty, and device condition so equipment does not get lost and support can be handled faster.",
        ),
        (
            "Software and Licence Tracking",
            [
                "Track software subscriptions, licence counts, assigned users, renewal dates, and vendor details.",
                "Identify unused or duplicate licences to reduce cost.",
                "Support audits by showing who has access to paid tools and why.",
            ],
            "For software assets, I track licences, subscriptions, renewal dates, and user assignments. This helps avoid paying for unused licences and makes it easier to prepare for renewals or audits.",
        ),
        (
            "Device Lifecycle",
            [
                "Follow the device from purchase to deployment, support, repair, recovery, replacement, and disposal.",
                "Know whether a device is in stock, assigned, under repair, retired, or missing.",
                "Plan replacements before old devices become a business problem.",
            ],
            "I think about devices as having a lifecycle. A laptop is purchased, prepared, assigned, supported, recovered, repaired if needed, and eventually replaced or disposed of securely. Tracking each stage helps the business avoid confusion and downtime.",
        ),
        (
            "Onboarding Support",
            [
                "Prepare devices before a new starter joins.",
                "Install required applications, configure accounts, assign licences, and record the asset.",
                "Make sure the user has the right tools from day one.",
            ],
            "Asset management makes onboarding smoother because IT can prepare the right device and software before the employee starts. I would make sure the device is configured, recorded in the register, assigned to the user, and ready with the required applications and access.",
        ),
        (
            "Offboarding and Recovery",
            [
                "Recover laptops, mobiles, access cards, chargers, and other company equipment.",
                "Update the asset register when equipment is returned.",
                "Check the device condition and decide whether it should be reused, repaired, wiped, or replaced.",
            ],
            "During offboarding, asset management is very important. I make sure company equipment is recovered, access is removed, the asset register is updated, and devices are wiped or prepared before being reused. This protects company data and reduces equipment loss.",
        ),
        (
            "Security and Compliance",
            [
                "Know which devices are company-owned and who is responsible for them.",
                "Support encryption, patching, antivirus, MDM, and secure wiping.",
                "Reduce risk when devices are lost, stolen, or returned by staff.",
            ],
            "Asset management supports security because the business needs to know where its devices are and who has access to company data. If a device is lost or a user leaves, accurate records help IT act quickly, disable access, wipe devices, or investigate risk.",
        ),
        (
            "Cost Control",
            [
                "Avoid unnecessary purchases by knowing what equipment is already available.",
                "Track repairs, warranties, and replacement needs.",
                "Reduce wasted software spend by removing unused licences.",
            ],
            "Good asset management helps control cost. If the register is accurate, the business can reuse available devices, claim warranty repairs, avoid duplicate purchases, and remove unused software licences.",
        ),
    ]

    for heading, bullets, answer in sections:
        doc.add_heading(heading, level=2)
        for bullet in bullets:
            add_bullet(doc, bullet)
        add_answer_box(doc, "Interview Answer", answer)

    doc.add_heading("4. Common Interview Questions and Answers", level=1)
    qas = [
        (
            "What information should be included in an asset register?",
            "I would include asset tag, serial number, device type, model, assigned user, department, location, purchase date, warranty status, device condition, software assigned, and current status such as in stock, deployed, under repair, or retired."
        ),
        (
            "How do you handle a new laptop purchase?",
            "I would record the device in the asset register, capture the serial number and warranty details, apply an asset tag if used, prepare the device with required settings and software, assign it to the user, and update the register with the deployment details."
        ),
        (
            "What would you do when an employee leaves?",
            "I would follow the offboarding checklist, recover all company equipment, update the asset register, check the device condition, remove or disable access, and arrange secure wiping or redeployment of the device."
        ),
        (
            "How does asset management improve security?",
            "It helps the company know which devices exist, who has them, and whether they contain company data. If a device is lost or a user leaves, accurate records help IT respond quickly by disabling access, wiping the device, or checking exposure."
        ),
        (
            "How do you reduce software licence waste?",
            "I would review licence assignments regularly, compare active users with actual business need, remove licences from departed staff, and identify tools that are duplicated or no longer used before renewal dates."
        ),
        (
            "What would you do if an asset register is outdated?",
            "I would clean it step by step. First, I would compare known records with actual devices, users, MDM data, invoices, and vendor records. Then I would update missing fields, identify unknown or missing assets, and create a simple process so the register stays updated going forward."
        ),
    ]

    for question, answer in qas:
        doc.add_heading(question, level=2)
        doc.add_paragraph(answer)

    doc.add_heading("5. STAR Example Answer", level=1)
    doc.add_paragraph("Use this when the interviewer asks for a real example of asset management.")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    set_cell_text(headers[0], "STAR Part", bold=True, color="FFFFFF")
    set_cell_text(headers[1], "Plain-Language Example", bold=True, color="FFFFFF")
    shade_cell(headers[0], "1F4E79")
    shade_cell(headers[1], "1F4E79")
    rows = [
        ("Situation", "The business needed a clearer view of IT equipment such as laptops, mobiles, printers, and accessories."),
        ("Task", "My responsibility was to keep the asset register accurate and make sure devices were tracked properly."),
        ("Action", "I recorded device details, assigned equipment to users, updated changes during onboarding and offboarding, tracked repairs, and followed up on returned or missing equipment."),
        ("Result", "The business had better visibility of IT assets, onboarding and offboarding became smoother, and equipment loss or unnecessary purchases were reduced."),
    ]
    for label, text in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], text)

    doc.add_heading("6. Words to Use in the Interview", level=1)
    phrases = [
        "I track assets across the full lifecycle.",
        "I keep the asset register updated when devices are purchased, assigned, repaired, returned, or retired.",
        "Asset management supports both cost control and security.",
        "During onboarding, accurate asset records help IT prepare devices faster.",
        "During offboarding, asset records help make sure company equipment and data are recovered securely.",
        "I try to keep licence records accurate so the business does not pay for unused software.",
        "I treat the asset register as a live operational record, not a one-time spreadsheet.",
    ]
    for phrase in phrases:
        add_bullet(doc, phrase)

    doc.add_heading("7. Final Polished Answer to Memorise", level=1)
    add_answer_box(
        doc,
        "Final Answer",
        "I have experience managing hardware and software assets, including laptops, tablets, mobiles, printers, accessories, and software licences. "
        "I maintain asset registers with details such as serial numbers, assigned users, locations, warranty status, device condition, and lifecycle status. "
        "I track devices from purchase through deployment, support, recovery, repair, replacement, and disposal. "
        "This helps the business control cost, avoid unnecessary purchases, recover equipment during offboarding, and protect company data. "
        "I see asset management as an important part of IT operations because it supports security, budgeting, onboarding, offboarding, and day-to-day support."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "Device and Asset Management Interview Preparation - Plain Language Guide"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
